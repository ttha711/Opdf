import json
import os
import re
import shutil
import subprocess
import sys
import base64
import time
import urllib.error
import urllib.request
from pathlib import Path
from difflib import SequenceMatcher
import io

import pdfplumber
from docx import Document
from docx.shared import Pt
from PIL import Image


def normalize_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def log_openrouter_usage(body: dict, tag: str) -> None:
    if not isinstance(body, dict):
        return
    usage = body.get("usage")
    if usage is None:
        return
    try:
        print(f"[openrouter_usage] {tag} {json.dumps(usage, ensure_ascii=False)}", file=sys.stderr)
    except Exception:
        print(f"[openrouter_usage] {tag} {usage}", file=sys.stderr)


def extract_blocks(pdf_path: Path):
    max_pages = int(os.environ.get("OPDF_REWRITE_MAX_PAGES", "0") or "0")
    pages = []
    with pdfplumber.open(str(pdf_path)) as pdf:
        for idx, page in enumerate(pdf.pages, start=1):
            if max_pages > 0 and idx > max_pages:
                break
            text = page.extract_text(layout=True) or ""
            raw_lines = [normalize_text(ln) for ln in text.splitlines()]
            lines = [ln for ln in raw_lines if ln]
            pages.append({"page": idx, "lines": lines})
    return pages


def extract_blocks_with_marker(pdf_path: Path):
    if os.environ.get("OPDF_USE_MARKER_EXTRACT", "0").strip() != "1":
        return None
    marker_exe = Path(__file__).resolve().parent / "python" / "runtime" / "Scripts" / "marker_single.exe"
    if not marker_exe.exists():
        return None
    max_pages = int(os.environ.get("OPDF_REWRITE_MAX_PAGES", "0") or "0")
    out_dir = Path(os.environ.get("OPDF_MARKER_TMP_DIR", str(Path.cwd() / ".tmp-marker"))).resolve()
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [str(marker_exe), str(pdf_path), "--output_format", "json", "--output_dir", str(out_dir), "--disable_multiprocessing", "--disable_tqdm"]
    if max_pages > 0:
        cmd.extend(["--page_range", ",".join(str(i) for i in range(max_pages))])
    subprocess.run(cmd, check=True, capture_output=True, text=True)
    json_file = out_dir / f"{pdf_path.stem}.json"
    if not json_file.exists():
        return None
    data = json.loads(json_file.read_text(encoding="utf-8"))
    pages = []
    # Marker JSON is markdown-centric, keep simple robust parsing.
    md = data.get("markdown", "") if isinstance(data, dict) else ""
    if not isinstance(md, str) or not md.strip():
        return None
    chunks = [c.strip() for c in md.split("\n\n") if c.strip()]
    # Rebuild pseudo pages for rewrite path.
    page_no = 1
    page_lines = []
    for c in chunks:
        lines = [normalize_text(x) for x in c.splitlines() if normalize_text(x)]
        if not lines:
            continue
        page_lines.extend(lines)
        if len(page_lines) >= 60:
            pages.append({"page": page_no, "lines": page_lines})
            page_no += 1
            page_lines = []
    if page_lines:
        pages.append({"page": page_no, "lines": page_lines})
    return pages if pages else None


def heuristic_rewrite_line(line: str, strict: bool) -> str:
    if strict:
        line = re.sub(r"\s*\[\d+\]", "", line)
        line = re.sub(r"\s{2,}", " ", line)
        return line.strip()
    line = re.sub(r"(?<=[A-Za-z])(?=[A-Z][a-z])", " ", line)
    line = re.sub(r"\s*\[\d+\]", "", line)
    line = re.sub(r"\s{2,}", " ", line)
    return line.strip()


def call_dify_rewrite(lines, strict: bool):
    dify_url = os.environ.get("OPDF_DIFY_URL", "").strip().rstrip("/")
    dify_key = os.environ.get("OPDF_DIFY_KEY", "").strip()
    ai_mode = os.environ.get("OPDF_AI_MODE", "local").strip().lower()
    if ai_mode != "dify" or not dify_url or not dify_key:
        return None

    prompt_payload = {
        "task": "ocr_rewrite",
        "strict": strict,
        "return_format": {"lines": "string[]", "same_length_as_input": True},
        "rules": [
            "Keep factual meaning exactly.",
            "Do not invent or summarize.",
            "Fix OCR typos, spacing, broken words.",
            "Return JSON object only: {\"lines\": [...]}"
        ],
        "lines": lines,
    }

    req = urllib.request.Request(
        f"{dify_url}/chat-messages",
        data=json.dumps(
            {
                "inputs": {},
                "query": json.dumps(prompt_payload, ensure_ascii=False),
                "response_mode": "blocking",
                "user": "opdf-desktop-converter",
            }
        ).encode("utf-8"),
        headers={"Authorization": f"Bearer {dify_key}", "Content-Type": "application/json"},
        method="POST",
    )

    timeout_sec = float(os.environ.get("OPDF_DIFY_TIMEOUT_SEC", "25"))
    with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
        body = json.loads(resp.read().decode("utf-8"))

    answer = (body.get("answer") or "").strip()
    if not answer:
        return None

    try:
        parsed = json.loads(answer)
        rewritten = parsed.get("lines") if isinstance(parsed, dict) else None
        if isinstance(rewritten, list) and len(rewritten) == len(lines):
            return [str(x).strip() for x in rewritten]
    except Exception:
        pass
    return None


def call_openai_rewrite(lines, strict: bool):
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None

    model = os.environ.get("OPDF_REWRITE_MODEL", "gpt-4.1-mini")
    system = (
        "You clean OCR text from PDFs. Preserve facts exactly. Do not add new facts. "
        "Keep original language. Return JSON array of strings with same length as input."
    )
    user = {
        "strict": strict,
        "instructions": "Fix OCR typos, spacing, broken words. Keep meaning and factual content unchanged.",
        "lines": lines,
    }

    payload = {
        "model": model,
        "input": [
            {"role": "system", "content": [{"type": "input_text", "text": system}]},
            {"role": "user", "content": [{"type": "input_text", "text": json.dumps(user, ensure_ascii=False)}]},
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "rewritten_lines",
                "schema": {
                    "type": "object",
                    "properties": {"lines": {"type": "array", "items": {"type": "string"}}},
                    "required": ["lines"],
                    "additionalProperties": False,
                },
            }
        },
    }

    req = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )

    with urllib.request.urlopen(req, timeout=90) as resp:
        body = json.loads(resp.read().decode("utf-8"))

    out_text = body.get("output_text", "")
    if not out_text:
        return None

    parsed = json.loads(out_text)
    rewritten = parsed.get("lines") if isinstance(parsed, dict) else None
    if not isinstance(rewritten, list) or len(rewritten) != len(lines):
        return None
    return [str(x).strip() for x in rewritten]


def call_openai_direct_pdf_rewrite(pdf_path: Path, strict: bool):
    api_key = os.environ.get("OPENAI_API_KEY", "").strip()
    if not api_key:
        return None
    if os.environ.get("OPDF_DIRECT_PDF_ENABLE", "0").strip() != "1":
        return None

    model = os.environ.get("OPDF_DIRECT_PDF_MODEL", "gpt-5")
    timeout_sec = float(os.environ.get("OPDF_DIRECT_PDF_TIMEOUT_SEC", "90"))

    with open(pdf_path, "rb") as f:
        pdf_b64 = base64.b64encode(f.read()).decode("utf-8")

    prompt = {
        "task": "rewrite_pdf_text_faithfully",
        "strict": strict,
        "rules": [
            "Preserve factual meaning exactly.",
            "Do not invent, summarize, or remove key facts.",
            "Fix OCR/spacing/broken words when needed.",
            "Return JSON object only with page-level lines."
        ],
        "format": {
            "pages": [{"page": "number", "lines": ["string"]}]
        },
    }

    payload = {
        "model": model,
        "input": [
            {
                "role": "user",
                "content": [
                    {"type": "input_file", "filename": pdf_path.name, "file_data": f"data:application/pdf;base64,{pdf_b64}"},
                    {"type": "input_text", "text": json.dumps(prompt, ensure_ascii=False)},
                ],
            }
        ],
        "text": {
            "format": {
                "type": "json_schema",
                "name": "rewritten_pdf_pages",
                "schema": {
                    "type": "object",
                    "properties": {
                        "pages": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "page": {"type": "integer"},
                                    "lines": {"type": "array", "items": {"type": "string"}},
                                },
                                "required": ["page", "lines"],
                                "additionalProperties": False,
                            },
                        }
                    },
                    "required": ["pages"],
                    "additionalProperties": False,
                },
            }
        },
    }

    req = urllib.request.Request(
        "https://api.openai.com/v1/responses",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    out_text = body.get("output_text", "")
    if not out_text:
        return None
    parsed = json.loads(out_text)
    pages = parsed.get("pages") if isinstance(parsed, dict) else None
    if not isinstance(pages, list):
        return None
    clean = []
    for p in pages:
        if not isinstance(p, dict):
            continue
        page_no = p.get("page")
        lines = p.get("lines")
        if isinstance(page_no, int) and isinstance(lines, list):
            clean.append({"page": page_no, "lines": [str(x).strip() for x in lines if str(x).strip()]})
    return clean if clean else None


def call_openrouter_direct_pdf_rewrite(pdf_path: Path, strict: bool):
    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not api_key:
        return None
    if os.environ.get("OPDF_DIRECT_PDF_ENABLE", "0").strip() != "1":
        return None

    model = os.environ.get("OPDF_OPENROUTER_MODEL", "google/gemini-2.5-flash")
    timeout_sec = float(os.environ.get("OPDF_DIRECT_PDF_TIMEOUT_SEC", "90"))
    debug = os.environ.get("OPDF_REWRITE_DEBUG_TIMING", "0").strip() == "1"

    t0 = time.perf_counter()
    with open(pdf_path, "rb") as f:
        pdf_bytes = f.read()
        pdf_b64 = base64.b64encode(pdf_bytes).decode("utf-8")
    t1 = time.perf_counter()
    if debug:
        print(f"[timing] openrouter.read_and_b64 sec={t1 - t0:.3f} bytes={len(pdf_bytes)}", file=sys.stderr)

    prompt = (
        "Rewrite this PDF text faithfully. Preserve factual meaning exactly, "
        "do not invent/summarize. Fix OCR/spacing/broken words when needed. "
        "Return JSON only with schema: {\"pages\":[{\"page\":number,\"lines\":string[]}]}."
    )

    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "file", "file": {"filename": pdf_path.name, "file_data": f"data:application/pdf;base64,{pdf_b64}"}},
                    {"type": "text", "text": json.dumps({"strict": strict}, ensure_ascii=False)},
                ],
            }
        ],
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "rewritten_pdf_pages",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "pages": {
                            "type": "array",
                            "items": {
                                "type": "object",
                                "properties": {
                                    "page": {"type": "integer"},
                                    "lines": {"type": "array", "items": {"type": "string"}},
                                },
                                "required": ["page", "lines"],
                                "additionalProperties": False,
                            },
                        }
                    },
                    "required": ["pages"],
                    "additionalProperties": False,
                },
            },
        },
    }

    t2 = time.perf_counter()
    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost",
            "X-Title": "OPDF",
        },
        method="POST",
    )
    t3 = time.perf_counter()
    if debug:
        print(f"[timing] openrouter.build_request sec={t3 - t2:.3f}", file=sys.stderr)
    t4 = time.perf_counter()
    with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    log_openrouter_usage(body, "direct_pdf_rewrite")
    t5 = time.perf_counter()
    if debug:
        print(f"[timing] openrouter.http_and_parse sec={t5 - t4:.3f}", file=sys.stderr)

    choices = body.get("choices") if isinstance(body, dict) else None
    if not isinstance(choices, list) or not choices:
        return None
    msg = (choices[0] or {}).get("message") or {}
    content = msg.get("content")
    if isinstance(content, list):
        text_parts = [x.get("text", "") for x in content if isinstance(x, dict)]
        out_text = "".join(text_parts).strip()
    else:
        out_text = (content or "").strip()
    if not out_text:
        return None

    parsed = json.loads(out_text)
    pages = parsed.get("pages") if isinstance(parsed, dict) else None
    if not isinstance(pages, list):
        return None
    clean = []
    for p in pages:
        if not isinstance(p, dict):
            continue
        page_no = p.get("page")
        lines = p.get("lines")
        if isinstance(page_no, int) and isinstance(lines, list):
            clean.append({"page": page_no, "lines": [str(x).strip() for x in lines if str(x).strip()]})
    return clean if clean else None


def call_openrouter_page_rewrite(lines, strict: bool):
    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not api_key:
        return None
    model = os.environ.get("OPDF_OPENROUTER_MODEL", "google/gemini-2.5-flash")
    timeout_sec = float(os.environ.get("OPDF_DIRECT_PDF_TIMEOUT_SEC", "90"))

    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Rewrite OCR text faithfully. Keep meaning exactly. "
                            "Do not invent new facts. Keep same number of lines."
                        ),
                    },
                    {
                        "type": "text",
                        "text": json.dumps(
                            {"strict": strict, "lines": lines},
                            ensure_ascii=False,
                        ),
                    },
                ],
            }
        ],
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "rewritten_lines",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {
                        "lines": {"type": "array", "items": {"type": "string"}}
                    },
                    "required": ["lines"],
                    "additionalProperties": False,
                },
            },
        },
    }

    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost",
            "X-Title": "OPDF",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    log_openrouter_usage(body, "page_rewrite")
    choices = body.get("choices") if isinstance(body, dict) else None
    if not isinstance(choices, list) or not choices:
        return None
    msg = (choices[0] or {}).get("message") or {}
    out_text = (msg.get("content") or "").strip()
    if not out_text:
        return None
    parsed = json.loads(out_text)
    rewritten = parsed.get("lines") if isinstance(parsed, dict) else None
    if not isinstance(rewritten, list) or len(rewritten) != len(lines):
        return None
    return [str(x).strip() for x in rewritten]


def call_openrouter_page_vision_rewrite(image_b64: str, strict: bool):
    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not api_key:
        return None
    model = os.environ.get("OPDF_OPENROUTER_MODEL", "google/gemini-2.5-flash")
    timeout_sec = float(os.environ.get("OPDF_DIRECT_PDF_TIMEOUT_SEC", "90"))
    payload = {
        "model": model,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "Extract and faithfully rewrite text from this document page image. Preserve meaning and order. Return JSON only: {\"lines\": string[]}"},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                    {"type": "text", "text": json.dumps({"strict": strict}, ensure_ascii=False)},
                ],
            }
        ],
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "vision_page_lines",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {"lines": {"type": "array", "items": {"type": "string"}}},
                    "required": ["lines"],
                    "additionalProperties": False,
                },
            },
        },
    }
    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://localhost",
            "X-Title": "OPDF",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=timeout_sec) as resp:
        body = json.loads(resp.read().decode("utf-8"))
    log_openrouter_usage(body, "vision_page_rewrite")
    choices = body.get("choices") if isinstance(body, dict) else None
    if not isinstance(choices, list) or not choices:
        return None
    content = ((choices[0] or {}).get("message") or {}).get("content")
    out_text = (content or "").strip() if isinstance(content, str) else ""
    if not out_text:
        return None
    parsed = json.loads(out_text)
    lines = parsed.get("lines") if isinstance(parsed, dict) else None
    if not isinstance(lines, list):
        return None
    return [str(x).strip() for x in lines if str(x).strip()]


def render_page_images(pdf_path: Path, page_numbers):
    images = {}
    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_no in page_numbers:
            if page_no < 1 or page_no > len(pdf.pages):
                continue
            pil = pdf.pages[page_no - 1].to_image(resolution=150).original.convert("RGB")
            buf = io.BytesIO()
            pil.save(buf, format="PNG", optimize=True)
            images[page_no] = base64.b64encode(buf.getvalue()).decode("utf-8")
    return images


def rewrite_blocks(pages, mode: str, source_pdf_path: Path | None = None):
    strict = mode == "strict-factual"
    out = []
    rewrite_source = "local"
    page_rewrite_details = []

    direct_enable = os.environ.get("OPDF_DIRECT_PDF_ENABLE", "0").strip() == "1"
    direct_mode = os.environ.get("OPDF_DIRECT_PDF_MODE", "full").strip().lower()
    min_similarity = float(os.environ.get("OPDF_REWRITE_MIN_SIMILARITY", "0.78"))

    if source_pdf_path is not None and direct_enable and direct_mode == "full":
        try:
            direct_pages = call_openrouter_direct_pdf_rewrite(source_pdf_path, strict)
            if direct_pages:
                rewrite_source = "openrouter_pdf_direct"
                direct_map = {p["page"]: p["lines"] for p in direct_pages}
                for page in pages:
                    lines = direct_map.get(page["page"], page["lines"])
                    out.append({"page": page["page"], "lines": lines})
                    page_rewrite_details.append({"page": page["page"], "source": "openrouter_pdf_direct", "fallback_reason": None})
                used_llm = True
                return out, used_llm, rewrite_source, page_rewrite_details
        except Exception:
            pass

        try:
            direct_pages = call_openai_direct_pdf_rewrite(source_pdf_path, strict)
            if direct_pages:
                rewrite_source = "openai_pdf_direct"
                direct_map = {p["page"]: p["lines"] for p in direct_pages}
                for page in pages:
                    lines = direct_map.get(page["page"], page["lines"])
                    out.append({"page": page["page"], "lines": lines})
                    page_rewrite_details.append({"page": page["page"], "source": "openai_pdf_direct", "fallback_reason": None})
                used_llm = True
                return out, used_llm, rewrite_source, page_rewrite_details
        except Exception:
            pass

    for page in pages:
        base_lines = [heuristic_rewrite_line(ln, strict) for ln in page["lines"]]
        base_lines = [ln for ln in base_lines if ln]
        rewritten = None
        page_source = "local"
        fallback_reason = None

        if base_lines and direct_enable:
            try:
                rewritten = call_openrouter_page_rewrite(base_lines, strict)
                if rewritten:
                    before = "\n".join(base_lines)
                    after = "\n".join(rewritten)
                    sim = SequenceMatcher(None, before, after).ratio()
                    if sim >= min_similarity:
                        rewrite_source = "openrouter_line_rewrite"
                        page_source = "openrouter_line_rewrite"
                    else:
                        rewritten = None
                        fallback_reason = f"openrouter_similarity_too_low:{sim:.3f}"
            except Exception:
                rewritten = None

        if base_lines:
            try:
                rewritten = call_dify_rewrite(base_lines, strict)
                if rewritten:
                    rewrite_source = "dify"
                    page_source = "dify"
                else:
                    fallback_reason = "dify_no_response_or_invalid_payload"
            except TimeoutError:
                fallback_reason = "dify_timeout"
                rewritten = None
            except urllib.error.URLError as exc:
                fallback_reason = f"dify_network_error:{exc.reason}"
                rewritten = None
            except Exception:
                fallback_reason = "dify_unknown_error"
                rewritten = None

        if not rewritten and base_lines:
            try:
                rewritten = call_openai_rewrite(base_lines, strict)
                if rewritten and rewrite_source != "dify":
                    rewrite_source = "openai"
                if rewritten:
                    page_source = "openai"
            except Exception:
                rewritten = None

        final_lines = rewritten if rewritten else base_lines
        out.append({"page": page["page"], "lines": final_lines})
        page_rewrite_details.append(
            {
                "page": page["page"],
                "source": page_source,
                "fallback_reason": fallback_reason if page_source != "dify" else None,
            }
        )

    used_llm = rewrite_source in {"dify", "openai"}

    # Vision fallback for scanned pages with empty text layer.
    if direct_enable and source_pdf_path is not None and all(len((x["lines"])) == 0 for x in out):
        try:
            page_numbers = [x["page"] for x in out]
            img_map = render_page_images(source_pdf_path, page_numbers)
            vision_out = []
            details = []
            used_any = False
            for page_no in page_numbers:
                lines = call_openrouter_page_vision_rewrite(img_map.get(page_no, ""), strict) if img_map.get(page_no) else None
                if lines:
                    used_any = True
                    vision_out.append({"page": page_no, "lines": lines})
                    details.append({"page": page_no, "source": "openrouter_vision_page", "fallback_reason": None})
                else:
                    vision_out.append({"page": page_no, "lines": []})
                    details.append({"page": page_no, "source": "local", "fallback_reason": "vision_no_output"})
            if used_any:
                return vision_out, True, "openrouter_vision_page", details
        except Exception:
            pass

    return out, used_llm, rewrite_source, page_rewrite_details


def compose_docx(pages, out_path: Path, source_name: str):
    doc = Document()
    doc.add_heading(f"Rewritten from PDF: {source_name}", level=1)
    for page in pages:
        if page["page"] > 1:
            doc.add_page_break()
        doc.add_heading(f"Page {page['page']}", level=2)
        buf = []

        def flush_buf():
            nonlocal buf
            if buf:
                doc.add_paragraph(" ".join(buf))
            buf = []

        for line in page["lines"]:
            if len(line) <= 90 and line[:1].isupper() and (line.endswith(":") or line.isupper()):
                flush_buf()
                doc.add_heading(line.rstrip(":"), level=3)
                continue
            if line.startswith(("- ", "* ")):
                flush_buf()
                p = doc.add_paragraph(line[2:].strip())
                p.style = "List Bullet"
                continue
            buf.append(line)
            if len(buf) >= 3:
                flush_buf()
        flush_buf()
    doc.save(str(out_path))


def postformat_docx_for_libreoffice(docx_path: Path):
    doc = Document(str(docx_path))
    title_size = int(os.environ.get("OPDF_DOCX_TITLE_SIZE_PT", "20"))
    h2_size = int(os.environ.get("OPDF_DOCX_H2_SIZE_PT", "14"))
    body_size = int(os.environ.get("OPDF_DOCX_BODY_SIZE_PT", "11"))

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue
        # Normalize spacing for cleaner LibreOffice reflow.
        p.paragraph_format.space_after = 0
        p.paragraph_format.space_before = 0
        if i == 0:
            p.style = "Title"
            for r in p.runs:
                r.font.size = Pt(title_size)
            continue
        if text.lower().startswith("page "):
            p.style = "Heading 2"
            for r in p.runs:
                r.font.size = Pt(h2_size)
            continue
        if text.startswith(("- ", "* ")):
            p.style = "List Bullet"
        else:
            p.style = "Normal"
        for r in p.runs:
            r.font.size = Pt(body_size)
    doc.save(str(docx_path))


def resolve_soffice():
    explicit = os.environ.get("OPDF_SOFFICE_PATH", "").strip()
    if explicit and Path(explicit).exists():
        return explicit
    which = shutil.which("soffice")
    if which:
        return which
    candidates = [
        r"C:\Program Files\LibreOffice\program\soffice.exe",
        r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    raise FileNotFoundError("LibreOffice soffice.exe not found. Set OPDF_SOFFICE_PATH.")


def export_pdf_via_libreoffice(docx_path: Path, out_pdf: Path):
    soffice = resolve_soffice()
    out_dir = out_pdf.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    cmd = [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)]
    subprocess.run(cmd, check=True, capture_output=True, text=True)


def main():
    if len(sys.argv) < 4:
        print(json.dumps({"ok": False, "error": "Usage: script <input.pdf> <format> <output> [mode]"}))
        sys.exit(2)

    in_path = Path(sys.argv[1]).resolve()
    fmt = sys.argv[2].lower()
    out_path = Path(sys.argv[3]).resolve()
    mode = (sys.argv[4] if len(sys.argv) > 4 else os.environ.get("OPDF_REWRITE_MODE", "strict-factual")).lower()
    if mode not in ["ai-rewrite", "strict-factual"]:
        mode = "strict-factual"
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Prefer the HTML semantic pipeline as main when enabled.
    if os.environ.get("OPDF_MAIN_PIPELINE", "html").strip().lower() == "html":
        alt_script = Path(__file__).resolve().parent / "pdf_to_html_rewrite_pipeline.py"
        if alt_script.exists():
            try:
                max_pages = int(os.environ.get("OPDF_REWRITE_MAX_PAGES", "0") or "0")
                cmd = [sys.executable, str(alt_script), str(in_path)]
                if max_pages > 0:
                    cmd.extend(["--max-pages", str(max_pages)])
                cmd.extend(["--out-docx", str(out_path)])
                cmd.extend(["--out-json", str(out_path.with_suffix(".json"))])
                res = subprocess.run(cmd, check=True, capture_output=True, text=True)
                # Pass through the child JSON for caller compatibility.
                print(res.stdout.strip())
                return
            except Exception:
                # Fall back to legacy flow if HTML pipeline fails.
                pass

    try:
        if fmt != "docx":
            raise ValueError("This pipeline currently supports docx only.")

        t_extract_start = time.perf_counter()
        pages = extract_blocks_with_marker(in_path) or extract_blocks(in_path)
        t_extract_end = time.perf_counter()
        pages, used_llm, rewrite_source, page_rewrite_details = rewrite_blocks(pages, mode, in_path)
        t_rewrite_end = time.perf_counter()
        compose_docx(pages, out_path, in_path.name)
        if os.environ.get("OPDF_DOCX_POSTFORMAT_ENABLE", "1").strip() == "1":
            postformat_docx_for_libreoffice(out_path)
        t_compose_end = time.perf_counter()

        pdf_out = out_path.with_suffix(".pdf")
        libre_pdf = {"ok": False, "path": str(pdf_out), "error": "skipped"}
        try:
            export_pdf_via_libreoffice(out_path, pdf_out)
            libre_pdf = {"ok": True, "path": str(pdf_out)}
        except Exception as exc:
            libre_pdf = {"ok": False, "path": str(pdf_out), "error": str(exc)}

        print(
            json.dumps(
                {
                    "ok": True,
                    "output": str(out_path),
                    "mode": mode,
                    "used_llm": used_llm,
                    "rewrite_source": rewrite_source,
                    "page_rewrite_details": page_rewrite_details,
                    "timings_sec": {
                        "extract_blocks": round(t_extract_end - t_extract_start, 3),
                        "rewrite_blocks": round(t_rewrite_end - t_extract_end, 3),
                        "compose_docx": round(t_compose_end - t_rewrite_end, 3),
                    },
                    "libreoffice_pdf": libre_pdf,
                }
            )
        )
    except Exception as exc:
        print(json.dumps({"ok": False, "error": str(exc)}))
        sys.exit(1)


if __name__ == "__main__":
    main()
