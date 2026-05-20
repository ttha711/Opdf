import argparse
import base64
import json
import os
import re
import sys
import time
import urllib.error
import urllib.request
from dataclasses import dataclass
from html.parser import HTMLParser
from io import BytesIO
from pathlib import Path

import pdfplumber
from docx import Document
import layout_schema_tools as lst
import docx_render_tools as drt



@dataclass
class Usage:
    prompt_tokens: int = 0
    completion_tokens: int = 0
    total_tokens: int = 0


def add_usage(total: Usage, inc: Usage) -> None:
    total.prompt_tokens += inc.prompt_tokens
    total.completion_tokens += inc.completion_tokens
    total.total_tokens += inc.total_tokens


def normalize_text(s: str) -> str:
    s = s.replace("\u00a0", " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()


def extract_page_text(page) -> list[str]:
    text = page.extract_text(layout=True) or ""
    if not text.strip():
        text = page.extract_text() or ""
    lines = [normalize_text(x) for x in text.splitlines()]
    return [x for x in lines if x]


def render_page_image_b64(page) -> str:
    pil = page.to_image(resolution=160).original.convert("RGB")
    buf = BytesIO()
    pil.save(buf, format="PNG")
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def call_openrouter_vision(image_b64: str, full_page_mode: bool = False) -> tuple[list[str], Usage] | tuple[None, Usage]:
    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not api_key:
        return None, Usage()

    model = os.environ.get("OPDF_OPENROUTER_MODEL", "google/gemini-2.5-flash")
    max_tokens = int(os.environ.get("OPDF_OPENROUTER_MAX_TOKENS", "1200") or "1200")
    prompt = (
        "Extract text faithfully from this PDF page image. "
        "Return the FULL page text in reading order. "
        "Do not summarize. Do not omit bullet points or footnotes. "
        "Return JSON only: {\"lines\": [string]}"
    )
    if full_page_mode:
        prompt = (
            "You must capture as much visible page text as possible in reading order. "
            "Keep exact wording, numbers, and punctuation. "
            "Return JSON only with many lines if needed: {\"lines\": [string]}"
        )

    payload = {
        "model": model,
        "temperature": 0.1,
        "max_tokens": max_tokens,
        "messages": [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                ],
            }
        ],
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": "vision_page_lines",
                "strict": True,
                "schema": {
                    "type": "object",
                    "properties": {"lines": {"type": "array", "items": {"type": "string"}}},
                    "required": ["lines"],
                    "additionalProperties": False,
                },
            },
        },
    }

    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "X-Title": "OPDF",
        },
        method="POST",
    )

    try:
        with urllib.request.urlopen(req, timeout=90) as resp:
            body = json.loads(resp.read().decode("utf-8"))
    except urllib.error.HTTPError as e:
        detail = ""
        try:
            detail = e.read().decode("utf-8", errors="replace")
        except Exception:
            detail = str(e)
        raise RuntimeError(f"openrouter_layout_schema_http_{e.code}: {detail}") from e

    usage_obj = body.get("usage") or {}
    usage = Usage(
        prompt_tokens=int(usage_obj.get("prompt_tokens", 0) or 0),
        completion_tokens=int(usage_obj.get("completion_tokens", 0) or 0),
        total_tokens=int(usage_obj.get("total_tokens", 0) or 0),
    )

    choices = body.get("choices") or []
    if not choices:
        return None, usage
    content = ((choices[0].get("message") or {}).get("content") or "").strip()
    if not content:
        return None, usage

    try:
        parsed = json.loads(content)
        lines = parsed.get("lines") if isinstance(parsed, dict) else None
        if isinstance(lines, list):
            normalized = [normalize_text(str(x)) for x in lines if normalize_text(str(x))]
            return normalized, usage
    except Exception:
        return None, usage
    return None, usage


def call_openrouter_layout_schema(image_b64: str) -> tuple[dict, Usage] | tuple[None, Usage]:
    api_key = os.environ.get("OPENROUTER_API_KEY", "").strip()
    if not api_key:
        return None, Usage()
    model = os.environ.get("OPDF_OPENROUTER_MODEL", "google/gemini-2.5-flash")
    max_tokens = int(os.environ.get("OPDF_OPENROUTER_MAX_TOKENS", "1200") or "1200")
    payload = {
        "model": model,
        "temperature": 0.1,
        "max_tokens": max_tokens,
        "messages": [
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract full page content and structure as JSON. "
                            "Preserve wording. Do NOT summarize. "
                            "Return JSON only with schema: "
                            "{\"blocks\":[{\"type\":\"heading|paragraph|list|table\","
                            "\"level\":1-3 optional,\"text\":\"...\" optional,"
                            "\"spans\":[{\"text\":\"...\",\"bold\":bool,\"italic\":bool,\"underline\":bool}] optional for paragraph,"
                            "\"items\":[\"...\"] optional,"
                            "\"rows\":[[\"c1\",\"c2\",...], ...] optional}]}"
                        ),
                    },
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image_b64}"}},
                ],
            }
        ],
        "response_format": {"type": "json_object"},
    }
    req = urllib.request.Request(
        "https://openrouter.ai/api/v1/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "X-Title": "OPDF",
        },
        method="POST",
    )
    with urllib.request.urlopen(req, timeout=90) as resp:
        body = json.loads(resp.read().decode("utf-8"))

    usage_obj = body.get("usage") or {}
    usage = Usage(
        prompt_tokens=int(usage_obj.get("prompt_tokens", 0) or 0),
        completion_tokens=int(usage_obj.get("completion_tokens", 0) or 0),
        total_tokens=int(usage_obj.get("total_tokens", 0) or 0),
    )
    choices = body.get("choices") or []
    if not choices:
        return None, usage
    content = ((choices[0].get("message") or {}).get("content") or "").strip()
    if not content:
        return None, usage
    try:
        parsed = json.loads(content)
        if isinstance(parsed, dict):
            # tolerate wrappers and recover from model variance
            if isinstance(parsed.get("blocks"), list):
                return parsed, usage
            if isinstance(parsed.get("data"), dict) and isinstance(parsed["data"].get("blocks"), list):
                return parsed["data"], usage
    except Exception:
        m = re.search(r"\{[\s\S]*\}", content)
        if m:
            try:
                parsed = json.loads(m.group(0))
                if isinstance(parsed, dict) and isinstance(parsed.get("blocks"), list):
                    return parsed, usage
            except Exception:
                return None, usage
    return None, usage


def detect_table_rows(lines: list[str]) -> list[list[str]] | None:
    rows = []
    for line in lines:
        if re.search(r"\S\s{2,}\S", line):
            cols = [normalize_text(c) for c in re.split(r"\s{2,}", line) if normalize_text(c)]
            if len(cols) >= 2:
                rows.append(cols)
    if len(rows) < 2:
        return None
    width = max(len(r) for r in rows)
    for r in rows:
        while len(r) < width:
            r.append("")
    return rows


def detect_block_type(line: str) -> str:
    if re.match(r"^([\-\*\u2022]|\d+[\.)])\s+", line):
        return "li"
    if len(line) <= 90 and (line.isupper() or re.match(r"^\d+(\.\d+)*\s+", line) or line.endswith(":")):
        return "h2"
    return "p"


def is_noise_line(line: str) -> bool:
    l = line.strip().lower()
    if not l:
        return True
    if re.match(r"^\d{1,2}/\d{1,2}/\d{2,4},?\s+\d{1,2}:\d{2}\s*(am|pm)", l):
        return True
    if l.startswith("http://") or l.startswith("https://") or "wikipedia" in l and "/" in l:
        return True
    if re.match(r"^\d+/\d+$", l):
        return True
    return False


def preprocess_lines(lines: list[str]) -> list[str]:
    out = []
    for line in lines:
        t = normalize_text(line)
        if not t or is_noise_line(t):
            continue
        out.append(t)
    return out


def join_wrapped_lines(lines: list[str]) -> list[str]:
    if not lines:
        return lines
    out: list[str] = []
    buf = lines[0]
    for line in lines[1:]:
        prev = buf
        if (
            not re.search(r"[.!?:;]$", prev)
            and not re.match(r"^([\-\*\u2022]|\d+[\.)])\s+", line)
            and len(prev) > 30
            and len(line) > 10
        ):
            buf = f"{prev} {line}"
        else:
            out.append(buf)
            buf = line
    out.append(buf)
    return out


def detect_kv_pairs(lines: list[str]) -> tuple[list[tuple[str, str]], list[str]]:
    pairs: list[tuple[str, str]] = []
    remain: list[str] = []
    i = 0
    while i < len(lines):
        cur = lines[i]
        nxt = lines[i + 1] if i + 1 < len(lines) else None
        if nxt and len(cur) <= 40 and not re.search(r"[.!?]$", cur):
            # likely "Label" + "Value" style block from infobox/metadata
            if len(nxt) <= 120:
                pairs.append((cur, nxt))
                i += 2
                continue
        remain.append(cur)
        i += 1
    return pairs, remain


def lines_to_html(lines: list[str], page_no: int) -> str:
    cleaned = join_wrapped_lines(preprocess_lines(lines))
    table = detect_table_rows(cleaned)
    html = [f"<h2>Page {page_no}</h2>"]
    kv_pairs, cleaned = detect_kv_pairs(cleaned)
    if kv_pairs:
        html.append("<table>")
        html.append("<tr><td><b>Field</b></td><td><b>Value</b></td></tr>")
        for k, v in kv_pairs:
            html.append(f"<tr><td>{escape_html(k)}</td><td>{escape_html(v)}</td></tr>")
        html.append("</table>")

    if table:
        html.append("<table>")
        for row in table:
            html.append("<tr>" + "".join(f"<td>{escape_html(c)}</td>" for c in row) + "</tr>")
        html.append("</table>")
        return "\n".join(html)

    in_ul = False
    for line in cleaned:
        kind = detect_block_type(line)
        cleaned = re.sub(r"^([\-\*\u2022]|\d+[\.)])\s+", "", line)
        if kind == "li":
            if not in_ul:
                html.append("<ul>")
                in_ul = True
            html.append(f"<li>{escape_html(cleaned)}</li>")
            continue
        if in_ul:
            html.append("</ul>")
            in_ul = False
        if kind == "h2":
            html.append(f"<h2>{escape_html(cleaned)}</h2>")
        else:
            html.append(f"<p>{escape_html(cleaned)}</p>")
    if in_ul:
        html.append("</ul>")
    return "\n".join(html)


def schema_to_html(schema: dict, page_no: int) -> str:
    html = [f"<h2>Page {page_no}</h2>"]
    blocks = schema.get("blocks") if isinstance(schema, dict) else None
    if not isinstance(blocks, list):
        return "\n".join(html)
    for b in blocks:
        if not isinstance(b, dict):
            continue
        btype = str(b.get("type", "")).strip().lower()
        if btype == "heading":
            level = int(b.get("level", 2) or 2)
            level = 2 if level < 2 or level > 3 else level
            text = normalize_text(str(b.get("text", "")))
            if text:
                html.append(f"<h{level}>{escape_html(text)}</h{level}>")
        elif btype == "paragraph":
            text = normalize_text(str(b.get("text", "")))
            if text:
                html.append(f"<p>{escape_html(text)}</p>")
        elif btype == "list":
            items = b.get("items")
            if isinstance(items, list) and items:
                html.append("<ul>")
                for item in items:
                    t = normalize_text(str(item))
                    if t:
                        html.append(f"<li>{escape_html(t)}</li>")
                html.append("</ul>")
        elif btype == "table":
            rows = b.get("rows")
            if isinstance(rows, list) and rows:
                html.append("<table>")
                for row in rows:
                    if not isinstance(row, list):
                        continue
                    cols = [escape_html(normalize_text(str(c))) for c in row if normalize_text(str(c))]
                    if cols:
                        html.append("<tr>" + "".join(f"<td>{c}</td>" for c in cols) + "</tr>")
                html.append("</table>")
    return "\n".join(html)


def escape_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


class DocxHtmlParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_ul = False
        self.current_table = None
        self.current_row = None
        self.current_cell = None
        self.buf = ""
        self.tag_stack = []

    def handle_starttag(self, tag, attrs):
        self.flush_text()
        self.tag_stack.append(tag)
        if tag == "table":
            self.current_table = self.doc.add_table(rows=0, cols=0)
        elif tag == "tr" and self.current_table is not None:
            self.current_row = None
        elif tag == "td":
            self.current_cell = ""

    def handle_endtag(self, tag):
        self.flush_text()
        if tag == "li":
            pass
        elif tag == "tr" and self.current_table is not None and isinstance(self.current_row, list):
            cols = len(self.current_row)
            if cols > 0:
                current_cols = len(self.current_table.columns)
                if current_cols == 0:
                    for _ in range(cols):
                        self.current_table.add_column(2000000)
                elif cols > current_cols:
                    for _ in range(cols - current_cols):
                        self.current_table.add_column(2000000)
                row_cells = self.current_table.add_row().cells
                for i, val in enumerate(self.current_row):
                    row_cells[i].text = val
            self.current_row = None
        elif tag == "td" and self.current_table is not None:
            if self.current_row is None:
                self.current_row = []
            self.current_row.append(self.current_cell or "")
            self.current_cell = None

        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

    def handle_data(self, data):
        self.buf += data

    def flush_text(self):
        text = normalize_text(self.buf)
        self.buf = ""
        if not text:
            return

        current = self.tag_stack[-1] if self.tag_stack else ""
        if current == "td":
            self.current_cell = (self.current_cell or "") + (" " if self.current_cell else "") + text
        elif current == "h2":
            self.doc.add_heading(text, level=2)
        elif current == "p":
            self.doc.add_paragraph(text)
        elif current == "li":
            self.doc.add_paragraph(text, style="List Bullet")


def postformat_docx(out_docx: Path):
    doc = Document(str(out_docx))
    title_size = int(os.environ.get("OPDF_DOCX_TITLE_SIZE_PT", "18"))
    h2_size = int(os.environ.get("OPDF_DOCX_H2_SIZE_PT", "13"))
    body_size = int(os.environ.get("OPDF_DOCX_BODY_SIZE_PT", "11"))
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        if idx == 0:
            p.style = "Title"
            for r in p.runs:
                r.font.size = Pt(title_size)
            continue
        if txt.lower().startswith("page "):
            p.style = "Heading 2"
            for r in p.runs:
                r.font.size = Pt(h2_size)
            continue
        if p.style and p.style.name == "List Bullet":
            for r in p.runs:
                r.font.size = Pt(body_size)
            continue
        p.style = "Normal"
        for r in p.runs:
            r.font.size = Pt(body_size)
    doc.save(str(out_docx))


def html_to_docx(html_pages: list[str], out_docx: Path):
    doc = Document()
    doc.add_heading("Rewritten Document", level=1)
    for i, html in enumerate(html_pages, start=1):
        parser = DocxHtmlParser(doc)
        parser.feed(html)
        if i < len(html_pages):
            doc.add_page_break()
    doc.save(str(out_docx))
    if os.environ.get("OPDF_DOCX_POSTFORMAT_ENABLE", "1").strip() == "1":
        postformat_docx(out_docx)


def run_pipeline(pdf_path: Path, out_docx: Path, out_json: Path, max_pages: int):
    t0 = time.perf_counter()
    pages = []
    html_pages = []
    rewrite_sources = []
    page_stats = []
    usage_total = Usage()
    min_lines = int(os.environ.get("OPDF_MIN_LINES_PER_PAGE", "10") or "10")

    with pdfplumber.open(str(pdf_path)) as pdf:
        limit = min(max_pages, len(pdf.pages)) if max_pages > 0 else len(pdf.pages)
        for idx in range(limit):
            page_no = idx + 1
            page = pdf.pages[idx]
            lines = extract_page_text(page)
            source = "text_layer" if lines else "no_text_layer"

            if not lines:
                image_b64 = render_page_image_b64(page)
                layout_schema, layout_usage = call_openrouter_layout_schema(image_b64)
                add_usage(usage_total, layout_usage)
                if layout_schema:
                    schema_html = lst.schema_to_html(layout_schema, page_no)
                    html_pages.append(schema_html)
                    rewrite_sources.append({"page": page_no, "rewrite_source": "openrouter_layout_schema"})
                    lines = []
                    pages.append({"page": page_no, "lines": lines})
                    page_stats.append(
                        {
                            "page": page_no,
                            "raw_line_count": 0,
                            "cleaned_line_count": 0,
                            "cleaned_preview": ["layout_schema_used"],
                        }
                    )
                    continue

                vision_lines, usage = call_openrouter_vision(image_b64)
                add_usage(usage_total, usage)
                # Quality gate: if extraction is too short, retry once with stricter full-page prompt.
                if vision_lines and len(vision_lines) < min_lines:
                    vision_lines_retry, usage_retry = call_openrouter_vision(image_b64, full_page_mode=True)
                    add_usage(usage_total, usage_retry)
                    if vision_lines_retry and len(vision_lines_retry) >= len(vision_lines):
                        vision_lines = vision_lines_retry
                if vision_lines:
                    lines = vision_lines
                    source = "openrouter_vision"
                elif os.environ.get("OPENROUTER_API_KEY", "").strip():
                    source = "openrouter_vision_no_output"
                else:
                    source = "no_text_no_openrouter"

            raw_count = len(lines)
            cleaned_preview = join_wrapped_lines(preprocess_lines(lines))
            pages.append({"page": page_no, "lines": lines})
            rewrite_sources.append({"page": page_no, "rewrite_source": source})
            html_pages.append(lines_to_html(lines, page_no))
            page_stats.append(
                {
                    "page": page_no,
                    "raw_line_count": raw_count,
                    "cleaned_line_count": len(cleaned_preview),
                    "cleaned_preview": cleaned_preview[:6],
                }
            )

    t1 = time.perf_counter()
    drt.html_to_docx(html_pages, out_docx)
    t2 = time.perf_counter()

    report = {
        "ok": True,
        "input": str(pdf_path),
        "output_docx": str(out_docx),
        "pages": len(pages),
        "rewrite_sources": rewrite_sources,
        "timings_sec": {
            "extract_and_rewrite": round(t1 - t0, 3),
            "html_to_docx": round(t2 - t1, 3),
            "total": round(t2 - t0, 3),
        },
        "openrouter_usage": {
            "prompt_tokens": usage_total.prompt_tokens,
            "completion_tokens": usage_total.completion_tokens,
            "total_tokens": usage_total.total_tokens,
        },
        "page_stats": page_stats,
    }
    out_json.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    return report


def main():
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass
    parser = argparse.ArgumentParser(description="Benchmark PDF -> HTML -> DOCX pipeline")
    parser.add_argument("pdf", type=Path)
    parser.add_argument("--max-pages", type=int, default=3)
    parser.add_argument("--out-docx", type=Path, default=Path(".tmp-html-3p.docx"))
    parser.add_argument("--out-json", type=Path, default=Path(".tmp-html-3p.json"))
    args = parser.parse_args()

    result = run_pipeline(args.pdf.resolve(), args.out_docx.resolve(), args.out_json.resolve(), args.max_pages)
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()
