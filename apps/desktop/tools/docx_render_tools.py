from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx_style_tools import SpanStyle, add_styled_text


class DocxHtmlParser(HTMLParser):
    def __init__(self, doc: Document):
        super().__init__()
        self.doc = doc
        self.current_table = None
        self.current_row = None
        self.current_cell = None
        self.current_paragraph = None
        self.buf = ""
        self.tag_stack = []

    def handle_starttag(self, tag, attrs):
        self.flush_text()
        self.tag_stack.append(tag)
        if tag == "table":
            self.current_table = self.doc.add_table(rows=0, cols=0)
        elif tag == "tr" and self.current_table is not None:
            self.current_row = None
        elif tag == "td":
            self.current_cell = ""
        elif tag in {"p", "li", "h2", "h3"}:
            if tag == "li":
                self.current_paragraph = self.doc.add_paragraph(style="List Bullet")
            elif tag == "h2":
                self.current_paragraph = self.doc.add_heading(level=2)
            elif tag == "h3":
                self.current_paragraph = self.doc.add_heading(level=3)
            else:
                self.current_paragraph = self.doc.add_paragraph()

    def handle_endtag(self, tag):
        self.flush_text()
        if tag == "tr" and self.current_table is not None and isinstance(self.current_row, list):
            cols = len(self.current_row)
            if cols > 0:
                current_cols = len(self.current_table.columns)
                if current_cols == 0:
                    for _ in range(cols):
                        self.current_table.add_column(2000000)
                elif cols > current_cols:
                    for _ in range(cols - current_cols):
                        self.current_table.add_column(2000000)
                row_cells = self.current_table.add_row().cells
                for i, val in enumerate(self.current_row):
                    row_cells[i].text = val
            self.current_row = None
        elif tag == "td" and self.current_table is not None:
            if self.current_row is None:
                self.current_row = []
            self.current_row.append(self.current_cell or "")
            self.current_cell = None

        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()
        if tag in {"p", "li", "h2", "h3"}:
            self.current_paragraph = None

    def handle_data(self, data):
        self.buf += data

    def flush_text(self):
        text = " ".join(self.buf.replace("\u00a0", " ").split()).strip()
        self.buf = ""
        if not text:
            return
        current = self.tag_stack[-1] if self.tag_stack else ""
        if current == "td":
            self.current_cell = (self.current_cell or "") + (" " if self.current_cell else "") + text
        elif current in {"h2", "h3", "p", "li"} and self.current_paragraph is not None:
            style = SpanStyle(
                bold=("b" in self.tag_stack or "strong" in self.tag_stack),
                italic=("i" in self.tag_stack or "em" in self.tag_stack),
                underline=("u" in self.tag_stack),
            )
            add_styled_text(self.current_paragraph, text, style)


def postformat_docx(out_docx: Path):
    doc = Document(str(out_docx))
    title_size = 18
    h2_size = 13
    body_size = 11
    for idx, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt:
            continue
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        if idx == 0:
            p.style = "Title"
            for r in p.runs:
                r.font.size = Pt(title_size)
            continue
        if txt.lower().startswith("page "):
            p.style = "Heading 2"
            for r in p.runs:
                r.font.size = Pt(h2_size)
            continue
        if p.style and p.style.name == "List Bullet":
            for r in p.runs:
                r.font.size = Pt(body_size)
            continue
        p.style = "Normal"
        for r in p.runs:
            r.font.size = Pt(body_size)
    doc.save(str(out_docx))


def html_to_docx(html_pages: list[str], out_docx: Path):
    doc = Document()
    doc.add_heading("Rewritten Document", level=1)
    for i, html in enumerate(html_pages, start=1):
        parser = DocxHtmlParser(doc)
        parser.feed(html)
        if i < len(html_pages):
            doc.add_page_break()
    doc.save(str(out_docx))
    postformat_docx(out_docx)
